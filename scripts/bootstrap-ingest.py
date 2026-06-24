"""One-shot bootstrap ingest: SharePoint -> Azure AI Search.

Use this script ONCE to backfill an empty AI Search index from an existing
SharePoint document library. Run it from a VM inside the VNet (Search and
Foundry are PE-only). After backfill, the Function App at
src/functions/spo-ingest takes over for autonomous delta sync.

All Azure auth via the VM's managed identity. Configuration via env vars:

    SPO_SITE_ID            SharePoint site ID (format: host,siteGuid,webGuid)
    AZURE_SEARCH_ENDPOINT  https://<srch>.search.windows.net
    AZURE_SEARCH_INDEX     Index name (default: ragdocs)
    AZURE_FOUNDRY_ENDPOINT https://<foundry>.openai.azure.com/  (or cognitiveservices.azure.com)
    EMBEDDING_DEPLOYMENT   Embedding deployment name (default: text-embedding-3-small)
    DOC_INTEL_ENDPOINT     Document Intelligence endpoint (optional; required only if PDFs in source)

Install deps once:
    pip install azure-identity azure-ai-documentintelligence azure-search-documents openai httpx python-docx
"""
from __future__ import annotations

import io
import logging
import os
from typing import Iterable

import httpx
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from azure.identity import ManagedIdentityCredential
from azure.search.documents import SearchClient
from openai import AzureOpenAI

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("bootstrap-ingest")


def env(name: str, default: str | None = None, *, required: bool = False) -> str:
    val = os.environ.get(name, default)
    if required and not val:
        raise SystemExit(f"Missing required env var: {name}")
    return val or ""


SITE_ID = env("SPO_SITE_ID", required=True)
SEARCH_ENDPOINT = env("AZURE_SEARCH_ENDPOINT", required=True)
INDEX_NAME = env("AZURE_SEARCH_INDEX", "ragdocs")
FOUNDRY_ENDPOINT = env("AZURE_FOUNDRY_ENDPOINT", required=True)
EMBED_DEPLOYMENT = env("EMBEDDING_DEPLOYMENT", "text-embedding-3-small")
DOC_INTEL_ENDPOINT = env("DOC_INTEL_ENDPOINT")  # optional

CHUNK_SIZE_CHARS = 1800
CHUNK_OVERLAP = 200
EMBED_DIM = 1536

cred = ManagedIdentityCredential()


def graph_token() -> str:
    return cred.get_token("https://graph.microsoft.com/.default").token


def graph_get(path: str) -> dict:
    h = {"Authorization": f"Bearer {graph_token()}"}
    r = httpx.get(f"https://graph.microsoft.com/v1.0{path}", headers=h, timeout=60)
    r.raise_for_status()
    return r.json()


def graph_get_bytes(path: str) -> bytes:
    # Graph /content endpoints return 302 redirects to SharePoint download.aspx
    # with a tempauth token. follow_redirects=True is REQUIRED.
    h = {"Authorization": f"Bearer {graph_token()}"}
    r = httpx.get(
        f"https://graph.microsoft.com/v1.0{path}",
        headers=h,
        timeout=180,
        follow_redirects=True,
    )
    r.raise_for_status()
    if len(r.content) < 200 or r.content[:1] == b"<":
        log.warning("Suspicious bytes from %s: first200=%r", path, r.content[:200])
    return r.content


def list_drive_files() -> list[dict]:
    """Return all .docx and .pdf files in the default doc lib.

    Sites.Selected 'read' allows GET /sites/{id}/drive (singular) but NOT
    /sites/{id}/drives (enumeration). Use the singular form for the default
    document library.
    """
    log.info("site_id=%s", SITE_ID)
    drive = graph_get(f"/sites/{SITE_ID}/drive")
    drive_id = drive["id"]
    log.info("default drive id=%s name=%s", drive_id, drive.get("name"))

    items = graph_get(f"/drives/{drive_id}/root/children")["value"]
    docs = [i for i in items if i["name"].lower().endswith((".docx", ".pdf"))]
    log.info("Found %d files: %s", len(docs), [d["name"] for d in docs])
    for d in docs:
        d["_drive_id"] = drive_id
    return docs


def download(item: dict) -> bytes:
    return graph_get_bytes(f"/drives/{item['_drive_id']}/items/{item['id']}/content")


def extract_text(filename: str, data: bytes) -> str:
    """Extract text using python-docx for .docx, Document Intelligence for PDFs.

    Document Intelligence's prebuilt-layout v4 *does* support .docx, but
    python-docx is faster and free for the simple-paragraph case. PDFs and
    scans use DI.
    """
    name_lower = filename.lower()
    if name_lower.endswith(".docx"):
        if data[:4] != b"PK\x03\x04":
            head = data[:200].decode("utf-8", errors="replace")
            raise ValueError(f"Not a .docx (bad magic). First 200 bytes: {head!r}")
        import docx as pydocx  # python-docx

        log.info("python-docx on %s (%d bytes)", filename, len(data))
        d = pydocx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)

    if not DOC_INTEL_ENDPOINT:
        raise SystemExit("PDF found but DOC_INTEL_ENDPOINT not set")
    di = DocumentIntelligenceClient(endpoint=DOC_INTEL_ENDPOINT, credential=cred)
    log.info("DI prebuilt-layout on %s (%d bytes)", filename, len(data))
    poller = di.begin_analyze_document(
        model_id="prebuilt-layout",
        body=AnalyzeDocumentRequest(bytes_source=data),
    )
    result = poller.result()
    return result.content or ""


def chunk_text(text: str, size: int = CHUNK_SIZE_CHARS, overlap: int = CHUNK_OVERLAP) -> Iterable[str]:
    text = text.strip()
    if not text:
        return
    i = 0
    while i < len(text):
        end = min(i + size, len(text))
        if end < len(text):
            sp = text.rfind(" ", i, end)
            if sp > i + size // 2:
                end = sp
        yield text[i:end].strip()
        if end >= len(text):
            break
        i = max(0, end - overlap)


def make_openai_client() -> AzureOpenAI:
    def token_provider() -> str:
        return cred.get_token("https://cognitiveservices.azure.com/.default").token

    return AzureOpenAI(
        api_version="2024-10-21",
        azure_endpoint=FOUNDRY_ENDPOINT,
        azure_ad_token_provider=token_provider,
    )


def main() -> None:
    oai = make_openai_client()
    search = SearchClient(endpoint=SEARCH_ENDPOINT, index_name=INDEX_NAME, credential=cred)

    docs = list_drive_files()
    if not docs:
        log.warning("No files to ingest.")
        return

    total_chunks = 0
    total_pushed = 0
    for d in docs:
        name = d["name"]
        web_url = d.get("webUrl", "")
        try:
            data = download(d)
            content = extract_text(name, data)
            log.info("%s: extracted %d chars", name, len(content))
            chunks = list(chunk_text(content))
            log.info("%s: %d chunks", name, len(chunks))

            search_docs = []
            for i in range(0, len(chunks), 16):
                batch = chunks[i : i + 16]
                resp = oai.embeddings.create(model=EMBED_DEPLOYMENT, input=batch)
                vectors = [e.embedding for e in resp.data]
                for j, (text, vec) in enumerate(zip(batch, vectors)):
                    idx = i + j
                    stem = name.replace(".docx", "").replace(".pdf", "").replace(" ", "_").lower()
                    search_docs.append(
                        {
                            "id": f"{stem}-{idx}",
                            "sourceDoc": name,
                            "sourceUrl": web_url,
                            "chunkIndex": idx,
                            "title": name.rsplit(".", 1)[0].replace("_", " "),
                            "content": text,
                            "contentVector": vec,
                        }
                    )

            res = search.upload_documents(documents=search_docs)
            n_pushed = sum(1 for r in res if r.succeeded)
            log.info("%s: pushed %d/%d chunks", name, n_pushed, len(search_docs))
            total_chunks += len(search_docs)
            total_pushed += n_pushed
        except Exception as e:
            log.exception("FAILED on %s: %s", name, e)

    log.info("=== DONE: %d/%d chunks indexed across %d docs ===", total_pushed, total_chunks, len(docs))


if __name__ == "__main__":
    main()
